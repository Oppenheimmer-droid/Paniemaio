"""
Servicio de documentos.
Maneja upload, procesamiento y eliminación de documentos.
"""

import os
import uuid
import hashlib
from typing import Optional, List, Tuple
from datetime import datetime

from sqlalchemy import select, and_, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import Document, DocumentChunk, DocumentStatus, User
from app.core.config import settings
from app.rag.vector_store import retrieval_pipeline


class DocumentService:
    """Servicio para gestión de documentos."""
    
    def __init__(self, db: AsyncSession):
        self.db = db
    
    async def save_file(
        self,
        file_content: bytes,
        filename: str,
        tenant_id: str
    ) -> str:
        """
        Guarda un archivo en el sistema de archivos.
        
        Args:
            file_content: Contenido del archivo
            filename: Nombre original del archivo
            tenant_id: ID del tenant
            
        Returns:
            Ruta relativa del archivo guardado
        """
        # Crear directorio si no existe
        upload_dir = os.path.join(settings.UPLOAD_DIR, tenant_id)
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generar nombre único
        ext = os.path.splitext(filename)[1].lower()
        unique_name = f"{uuid.uuid4()}{ext}"
        
        # Guardar archivo
        file_path = os.path.join(upload_dir, unique_name)
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        # Devolver ruta relativa
        relative_path = os.path.join(tenant_id, unique_name)
        return relative_path
    
    async def create_document(
        self,
        tenant_id: str,
        user_id: str,
        filename: str,
        file_path: str,
        file_size: int,
        mime_type: str,
        title: str,
        subject_id: Optional[str] = None,
        topic_id: Optional[str] = None,
        description: Optional[str] = None,
        difficulty: int = 1
    ) -> Document:
        """
        Crea un nuevo registro de documento.
        
        Args:
            tenant_id: ID del tenant
            user_id: ID del usuario que sube
            filename: Nombre del archivo
            file_path: Ruta del archivo
            file_size: Tamaño en bytes
            mime_type: Tipo MIME
            title: Título del documento
            subject_id: ID de materia (opcional)
            topic_id: ID de tema (opcional)
            description: Descripción (opcional)
            difficulty: Dificultad (1-5)
            
        Returns:
            Document creado
        """
        document = Document(
            id=str(uuid.uuid4()),
            tenant_id=tenant_id,
            uploaded_by=user_id,
            filename=filename,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type,
            title=title,
            subject_id=subject_id,
            topic_id=topic_id,
            description=description,
            difficulty=difficulty,
            status=DocumentStatus.PENDING.value,
            page_count=0,
            chunk_count=0
        )
        
        self.db.add(document)
        await self.db.flush()
        
        return document
    
    async def get_document(
        self,
        document_id: str,
        tenant_id: str
    ) -> Optional[Document]:
        """Obtiene un documento por ID."""
        result = await self.db.execute(
            select(Document).where(
                and_(
                    Document.id == document_id,
                    Document.tenant_id == tenant_id
                )
            )
        )
        return result.scalar_one_or_none()
    
    async def list_documents(
        self,
        tenant_id: str,
        page: int = 1,
        page_size: int = 20,
        status: Optional[str] = None
    ) -> Tuple[List[Document], int]:
        """
        Lista documentos de un tenant con paginación.
        
        Returns:
            (documents, total_count)
        """
        # Query base
        query = select(Document).where(Document.tenant_id == tenant_id)
        count_query = select(func.count(Document.id)).where(Document.tenant_id == tenant_id)
        
        # Filtrar por status si se especifica
        if status:
            query = query.where(Document.status == status)
            count_query = count_query.where(Document.status == status)
        
        # Ordenar por fecha de creación (más reciente primero)
        query = query.order_by(Document.created_at.desc())
        
        # Aplicar paginación
        offset = (page - 1) * page_size
        query = query.offset(offset).limit(page_size)
        
        # Ejecutar queries
        result = await self.db.execute(query)
        documents = result.scalars().all()
        
        count_result = await self.db.execute(count_query)
        total = count_result.scalar()
        
        return list(documents), total
    
    async def update_status(
        self,
        document_id: str,
        status: str,
        error_message: Optional[str] = None,
        page_count: Optional[int] = None,
        chunk_count: Optional[int] = None
    ) -> bool:
        """Actualiza el estado de un documento."""
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return False
        
        document.status = status
        
        if error_message:
            document.error_message = error_message
        
        if page_count is not None:
            document.page_count = page_count
        
        if chunk_count is not None:
            document.chunk_count = chunk_count
        
        if status == DocumentStatus.COMPLETED.value:
            document.processed_at = datetime.utcnow()
        
        await self.db.flush()
        return True
    
    async def delete_document(
        self,
        document_id: str,
        tenant_id: str
    ) -> bool:
        """
        Elimina un documento y sus vectores asociados.
        
        Returns:
            True si se eliminó correctamente
        """
        document = await self.get_document(document_id, tenant_id)
        if not document:
            return False
        
        # Eliminar vectores de ChromaDB
        try:
            retrieval_pipeline.delete_document_vectors(tenant_id, document_id)
        except Exception as e:
            print(f"Error al eliminar vectores: {e}")
        
        # Eliminar archivo físico
        try:
            full_path = os.path.join(settings.UPLOAD_DIR, document.file_path)
            if os.path.exists(full_path):
                os.remove(full_path)
        except Exception as e:
            print(f"Error al eliminar archivo: {e}")
        
        # Eliminar chunks de la BD
        chunk_result = await self.db.execute(
            select(DocumentChunk).where(DocumentChunk.document_id == document_id)
        )
        chunks = chunk_result.scalars().all()
        for chunk in chunks:
            await self.db.delete(chunk)
        
        # Eliminar documento
        await self.db.delete(document)
        await self.db.flush()
        
        return True


class DocumentProcessor:
    """
    Procesador de documentos.
    Extrae texto, genera chunks y embeddings.
    """
    
    @staticmethod
    def extract_text(file_path: str, mime_type: str) -> Tuple[str, int]:
        """
        Extrae texto de un archivo.
        
        Args:
            file_path: Ruta al archivo
            mime_type: Tipo MIME del archivo
            
        Returns:
            (text, page_count)
        """
        full_path = os.path.join(settings.UPLOAD_DIR, file_path)
        
        if not os.path.exists(full_path):
            raise FileNotFoundError(f"Archivo no encontrado: {full_path}")
        
        text = ""
        page_count = 0
        
        if mime_type == "application/pdf":
            text, page_count = DocumentProcessor._extract_pdf(full_path)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            text, page_count = DocumentProcessor._extract_docx(full_path)
        else:
            # Intentar como texto plano
            try:
                with open(full_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except:
                raise ValueError(f"Tipo de archivo no soportado: {mime_type}")
        
        return text, page_count
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Tuple[str, int]:
        """Extrae texto de PDF."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n\n".join(text_parts), page_count
            
        except ImportError:
            raise ImportError("pypdf no está instalado")
    
    @staticmethod
    def _extract_docx(file_path: str) -> Tuple[str, int]:
        """Extrae texto de DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            page_count = 1  # DOCX no tiene páginas reales
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            return "\n\n".join(text_parts), page_count
            
        except ImportError:
            raise ImportError("python-docx no está instalado")
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = None,
        overlap: int = None
    ) -> List[Dict[str, Any]]:
        """
        Divide texto en chunks.
        
        Args:
            text: Texto a dividir
            chunk_size: Tamaño de cada chunk (default de settings)
            overlap: Superposición entre chunks (default de settings)
            
        Returns:
            Lista de dicts con {content, chunk_index, start_char, end_char}
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        chunks = []
        text_length = len(text)
        
        if text_length == 0:
            return chunks
        
        start = 0
        chunk_index = 0
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Si no es el último chunk, intentar cortar en espacio
            if end < text_length:
                # Buscar último espacio en el rango
                last_space = text.rfind(" ", start, end)
                if last_space > start:
                    end = last_space
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Solo añadir si hay contenido
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end
                })
            
            # Avanzar con overlap
            start = end - overlap
            chunk_index += 1
            
            if start <= 0:
                break
        
        return chunks
    
    @staticmethod
    async def process_document(
        document_id: str,
        tenant_id: str,
        db: AsyncSession
    ) -> Tuple[int, List[str]]:
        """
        Procesa un documento: extrae texto, genera chunks y embeddings.
        
        Args:
            document_id: ID del documento
            tenant_id: ID del tenant
            db: Sesión de base de datos
            
        Returns:
            (chunk_count, vector_ids)
        """
        # Obtener documento
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            raise ValueError(f"Documento no encontrado: {document_id}")
        
        # Actualizar estado a processing
        document.status = DocumentStatus.PROCESSING.value
        await db.flush()
        
        try:
            # 1. Extraer texto
            print(f"📄 Extrayendo texto de: {document.filename}")
            text, page_count = DocumentProcessor.extract_text(
                document.file_path,
                document.mime_type
            )
            
            # 2. Dividir en chunks
            print(f"✂️  Generando {len(text)} chunks...")
            chunks_data = DocumentProcessor.chunk_text(text)
            
            # 3. Preparar chunks con metadata
            chunks = []
            for chunk in chunks_data:
                chunks.append({
                    "content": chunk["content"],
                    "document_id": document_id,
                    "chunk_index": chunk["chunk_index"],
                    "page_number": None  # Se podría mejorar extrayendo páginas
                })
            
            # 4. Guardar chunks en BD
            print(f"💾 Guardando chunks en BD...")
            for chunk_data in chunks:
                db_chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    tenant_id=tenant_id,
                    content=chunk_data["content"],
                    chunk_index=chunk_data["chunk_index"],
                    page_number=chunk_data.get("page_number"),
                    start_char=chunk_data.get("start_char", 0),
                    end_char=chunk_data.get("end_char", 0)
                )
                db.add(db_chunk)
            
            await db.flush()
            
            # 5. Añadir a ChromaDB
            print(f"🔢 Añadiendo a ChromaDB...")
            chunk_count, vector_ids = retrieval_pipeline.add_documents(
                tenant_id, chunks
            )
            
            # 6. Actualizar documento
            document.status = DocumentStatus.COMPLETED.value
            document.page_count = page_count
            document.chunk_count = chunk_count
            document.processed_at = datetime.utcnow()
            await db.flush()
            
            print(f"✅ Documento procesado: {chunk_count} chunks")
            return chunk_count, vector_ids
            
        except Exception as e:
            print(f"❌ Error procesando documento: {e}")
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)
            await db.flush()
            raise